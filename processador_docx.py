import os
import re
import copy
import uuid
import shutil
import tempfile
import zipfile
import base64
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from lxml import etree

def remover_paginas_em_branco(caminho_docx: str):
    doc = Document(caminho_docx)
    body = doc.element.body
    elementos = list(body)
    
    for el in reversed(elementos):
        if el.tag.endswith('sectPr'):
            continue
        if el.tag.endswith('p'):
            p = Paragraph(el, doc._body)
            texto = p.text.strip()
            tem_quebra_hard = bool(el.xpath('.//w:br[@w:type="page"]'))
            tem_quebra_api = getattr(p, 'contains_page_break', False)
            tem_quebra_format = bool(p.paragraph_format and p.paragraph_format.page_break_before)
            
            if not texto and not (tem_quebra_hard or tem_quebra_api or tem_quebra_format):
                try:
                    parent = el.getparent()
                    if parent is not None:
                        ps = parent.findall(qn('w:p'))
                        if len(ps) > 1:
                            parent.remove(el)
                        else:
                            break
                except Exception:
                    pass
            else:
                break
        else:
            break

    elementos = list(body)
    ultimo_foi_quebra = True
    
    for el in elementos:
        if el.tag.endswith('p'):
            p = Paragraph(el, doc._body)
            texto = p.text.strip()
            tem_quebra_hard_els = el.xpath('.//w:br[@w:type="page"]')
            tem_quebra_api = getattr(p, 'contains_page_break', False)
            tem_quebra_format = bool(p.paragraph_format and p.paragraph_format.page_break_before)
            
            tem_quebra_atual = bool(tem_quebra_hard_els) or tem_quebra_api or tem_quebra_format
            
            if tem_quebra_atual:
                if ultimo_foi_quebra and not texto:
                    for br in tem_quebra_hard_els:
                        try:
                            br.getparent().remove(br)
                        except Exception:
                            pass
                    
                    if tem_quebra_format and p.paragraph_format:
                        p.paragraph_format.page_break_before = False
                        
                    if tem_quebra_api or getattr(p, 'rendered_page_breaks', []):
                        for rpb in el.xpath('.//w:lastRenderedPageBreak'):
                            try:
                                rpb.getparent().remove(rpb)
                            except Exception:
                                pass
                ultimo_foi_quebra = True
            elif texto:
                ultimo_foi_quebra = False
        elif el.tag.endswith('tbl'):
            ultimo_foi_quebra = False
            
    doc.save(caminho_docx)

def _remover_ultimo_sectpr(caminho_docx):
    try:
        with zipfile.ZipFile(caminho_docx, 'r') as zin:
            conteudos = {n: zin.read(n) for n in zin.namelist()}

        if 'word/document.xml' not in conteudos:
            return

        tree = etree.fromstring(conteudos['word/document.xml'])
        
        sect_prs_in_p = tree.findall(f'.//{qn("w:pPr")}/{qn("w:sectPr")}')
        if sect_prs_in_p:
            ultimo_sect_pr = sect_prs_in_p[-1]
            ultimo_sect_pr.getparent().remove(ultimo_sect_pr)

            conteudos['word/document.xml'] = etree.tostring(
                tree, xml_declaration=True, encoding='UTF-8', standalone=True
            )

            fd, temp_path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for nome, conteudo in conteudos.items():
                    zout.writestr(nome, conteudo)

            shutil.move(temp_path, caminho_docx)

    except Exception:
        pass

def _to_roman(n):
    try:
        n = int(n)
    except Exception:
        return "1"
    if n <= 0 or n >= 4000:
        return str(n)
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syb = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman_num = ""
    i = 0
    while n > 0:
        for _ in range(n // val[i]):
            roman_num += syb[i]
            n -= val[i]
        i += 1
    return roman_num

def _format_number(val, fmt):
    if fmt == 'decimal':
        return str(val)
    if fmt == 'lowerLetter':
        return chr(97 + (val - 1) % 26) if val > 0 else str(val)
    if fmt == 'upperLetter':
        return chr(65 + (val - 1) % 26) if val > 0 else str(val)
    if fmt == 'lowerRoman':
        return _to_roman(val).lower()
    if fmt == 'upperRoman':
        return _to_roman(val).upper()
    if fmt == 'decimalZero':
        return f"{val:02d}"
    if fmt == 'ordinal':
        return f"{val}º"
    if fmt in ('bullet', 'none'):
        return ""
    return str(val)

def _converter_listas_para_texto(caminho_docx):
    try:
        with zipfile.ZipFile(caminho_docx, 'r') as zin:
            nomes = zin.namelist()
            conteudos = {n: zin.read(n) for n in nomes}
    except Exception:
        return caminho_docx, False

    if 'word/numbering.xml' not in conteudos:
        return caminho_docx, False

    try:
        if 'word/styles.xml' in conteudos:
            styles_tree = etree.fromstring(conteudos['word/styles.xml'])
            modificado_styles = False
            for style in styles_tree.findall(qn('w:style')):
                numPrs = style.findall(f".//{qn('w:numPr')}")
                for numPr in numPrs:
                    numPr.getparent().remove(numPr)
                    modificado_styles = True
                numStyleLinks = style.findall(f".//{qn('w:numStyleLink')}")
                for numStyleLink in numStyleLinks:
                    numStyleLink.getparent().remove(numStyleLink)
                    modificado_styles = True
            if modificado_styles:
                conteudos['word/styles.xml'] = etree.tostring(styles_tree, xml_declaration=True, encoding='UTF-8', standalone=True)

        num_tree = etree.fromstring(conteudos['word/numbering.xml'])
        abstract_nums = {}
        for abs_num in num_tree.findall(qn('w:abstractNum')):
            abs_id = abs_num.get(qn('w:abstractNumId'))
            abstract_nums[abs_id] = {}
            for lvl in abs_num.findall(qn('w:lvl')):
                ilvl = int(lvl.get(qn('w:ilvl'), '0'))
                start_el = lvl.find(qn('w:start'))
                numFmt_el = lvl.find(qn('w:numFmt'))
                lvlText_el = lvl.find(qn('w:lvlText'))
                lvl_pPr = lvl.find(qn('w:pPr'))
                lvl_rPr = lvl.find(qn('w:rPr'))
                abstract_nums[abs_id][ilvl] = {
                    'start': int(start_el.get(qn('w:val'), '1')) if start_el is not None else 1,
                    'fmt': numFmt_el.get(qn('w:val')) if numFmt_el is not None else 'decimal',
                    'text': lvlText_el.get(qn('w:val')) if lvlText_el is not None else '',
                    'pPr': copy.deepcopy(lvl_pPr) if lvl_pPr is not None else None,
                    'rPr': copy.deepcopy(lvl_rPr) if lvl_rPr is not None else None,
                }

        nums = {}
        for num in num_tree.findall(qn('w:num')):
            num_id = num.get(qn('w:numId'))
            abs_ref = num.find(qn('w:abstractNumId'))
            abs_val = abs_ref.get(qn('w:val')) if abs_ref is not None else '0'
            nums[num_id] = {'abstractNumId': abs_val, 'overrides': {}}
            for lvl_override in num.findall(qn('w:lvlOverride')):
                ilvl = int(lvl_override.get(qn('w:ilvl'), '0'))
                start_ov = lvl_override.find(qn('w:startOverride'))
                if start_ov is not None:
                    nums[num_id]['overrides'][ilvl] = int(start_ov.get(qn('w:val'), '1'))

        style_num_map = {}
        if 'word/styles.xml' in conteudos:
            for style in styles_tree.findall(qn('w:style')):
                style_id = style.get(qn('w:styleId'))
                numPr = style.find(f".//{qn('w:numPr')}")
                if numPr is not None:
                    numId_el = numPr.find(qn('w:numId'))
                    ilvl_el = numPr.find(qn('w:ilvl'))
                    if numId_el is not None:
                        nId = numId_el.get(qn('w:val'))
                        ilvl = int(ilvl_el.get(qn('w:val'), '0')) if ilvl_el is not None else 0
                        style_num_map[style_id] = (nId, ilvl)

        xml_targets = [
            n for n in conteudos
            if n.startswith('word/') and n.endswith('.xml')
            and n not in ('word/numbering.xml', 'word/styles.xml',
                          'word/settings.xml', 'word/fontTable.xml', 'word/webSettings.xml')
        ]

        for target in xml_targets:
            try:
                tree = etree.fromstring(conteudos[target])
                counters = {}
                modificado = False
                for p in tree.iter(qn('w:p')):
                    nId = None
                    ilvl = 0
                    from_style = False
                    
                    numPr = p.find(f".//{qn('w:numPr')}")
                    if numPr is not None:
                        numId_el = numPr.find(qn('w:numId'))
                        ilvl_el = numPr.find(qn('w:ilvl'))
                        if numId_el is not None:
                            nId = numId_el.get(qn('w:val'))
                            if nId == '0':
                                nId = None
                            else:
                                ilvl = int(ilvl_el.get(qn('w:val'), '0')) if ilvl_el is not None else 0
                    else:
                        pStyle = p.find(f".//{qn('w:pStyle')}")
                        if pStyle is not None:
                            style_id = pStyle.get(qn('w:val'))
                            if style_id in style_num_map:
                                nId, ilvl = style_num_map[style_id]
                                from_style = True

                    if nId and nId in nums:
                        abs_id = nums[nId]['abstractNumId']
                        if abs_id in abstract_nums and ilvl in abstract_nums[abs_id]:
                            lvl_info = abstract_nums[abs_id][ilvl]
                            
                            if nId not in counters:
                                counters[nId] = {}
                                
                            if ilvl not in counters[nId]:
                                start_val = nums[nId]['overrides'].get(ilvl, lvl_info['start'])
                                counters[nId][ilvl] = start_val
                            else:
                                counters[nId][ilvl] += 1
                                
                            for l in list(counters[nId].keys()):
                                if l > ilvl:
                                    del counters[nId][l]
                                    
                            text_format = lvl_info['text']
                            for level in range(9):
                                placeholder = f"%{level+1}"
                                if placeholder in text_format:
                                    val = counters[nId].get(level, abstract_nums[abs_id].get(level, {}).get('start', 1))
                                    fmt = abstract_nums[abs_id].get(level, {}).get('fmt', 'decimal')
                                    text_format = text_format.replace(placeholder, _format_number(val, fmt))
                                    
                            run = OxmlElement('w:r')
                            t = OxmlElement('w:t')
                            t.text = text_format + " "
                            t.set(qn('xml:space'), 'preserve')
                            run.append(t)
                            
                            lvl_rPr = lvl_info.get('rPr')
                            if lvl_rPr is not None:
                                run.insert(0, copy.deepcopy(lvl_rPr))
                            else:
                                primeiro_run = p.find(qn('w:r'))
                                if primeiro_run is not None:
                                    rPr_primeiro = primeiro_run.find(qn('w:rPr'))
                                    if rPr_primeiro is not None:
                                        run.insert(0, copy.deepcopy(rPr_primeiro))
                                        
                            pPr = p.find(qn('w:pPr'))
                            if pPr is None:
                                pPr = OxmlElement('w:pPr')
                                p.insert(0, pPr)
                            else:
                                for el_remover in pPr.findall(qn('w:numPr')):
                                    el_remover.getparent().remove(el_remover)
                                    
                            lvl_pPr = lvl_info.get('pPr')
                            if lvl_pPr is not None:
                                lvl_ind = lvl_pPr.find(qn('w:ind'))
                                if lvl_ind is not None and pPr.find(qn('w:ind')) is None:
                                    pPr.append(copy.deepcopy(lvl_ind))
                                    
                            p.insert(list(p).index(pPr) + 1, run)
                            
                            if from_style:
                                new_numPr = OxmlElement('w:numPr')
                                new_numId = OxmlElement('w:numId')
                                new_numId.set(qn('w:val'), '0')
                                new_numPr.append(new_numId)
                                pPr.append(new_numPr)
                                
                            modificado = True

                if modificado:
                    conteudos[target] = etree.tostring(
                        tree, xml_declaration=True, encoding='UTF-8', standalone=True
                    )
            except Exception:
                pass

        conteudos['word/numbering.xml'] = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:numbering>'

        fd, temp_path = tempfile.mkstemp(suffix='.docx')
        os.close(fd)
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for nome, conteudo in conteudos.items():
                zout.writestr(nome, conteudo)

        return temp_path, True

    except Exception:
        return caminho_docx, False

def _mesclar_docx_no_local(doc, paragrafo_alvo, caminho_docx):
    arquivo_importar, eh_temp = _converter_listas_para_texto(caminho_docx)

    try:
        with open(arquivo_importar, 'rb') as f:
            docx_bytes = f.read()
    except Exception as e:
        _substituir_texto_mantendo_formatacao(paragrafo_alvo, paragrafo_alvo.text, f"[ERRO AO LER DOCX: {e}]")
        if eh_temp:
            try:
                os.unlink(arquivo_importar)
            except Exception:
                pass
        return

    part = doc.part
    chunk_name = f'/word/embeddings/chunk_{uuid.uuid4().hex}.docx'
    uri = PackURI(chunk_name)
    
    alt_part = Part(
        partname=uri,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        blob=docx_bytes,
        package=part.package
    )
    
    rel_id = part.relate_to(alt_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/aFChunk')
    
    altChunk = OxmlElement('w:altChunk')
    altChunk.set(qn('r:id'), rel_id)
    
    altChunkPr = OxmlElement('w:altChunkPr')
    matchSrc = OxmlElement('w:matchSrc')
    matchSrc.set(qn('w:val'), '1')
    altChunkPr.append(matchSrc)
    altChunk.append(altChunkPr)
    
    parent = paragrafo_alvo._element.getparent()
    if parent is not None:
        idx = list(parent).index(paragrafo_alvo._element)
        parent.insert(idx, altChunk)
        parent.remove(paragrafo_alvo._element)
        if len(parent.findall(qn('w:p'))) == 0:
            parent.append(OxmlElement('w:p'))

    if eh_temp:
        try:
            os.unlink(arquivo_importar)
        except Exception:
            pass

def _remover_secao_arp(doc):
    body = doc.element.body
    elementos_body = list(body)
    arp_element = None
    
    for p in reversed(doc.paragraphs):
        texto = p.text.strip().upper()
        if "MINUTA DA ATA DE REGISTRO DE PREÇOS" in texto or "MINUTA DE ATA DE REGISTRO DE PREÇOS" in texto:
            if len(texto) < 200:
                arp_element = p._element
                break
                
    for p in doc.paragraphs:
        texto = p.text.strip().upper()
        if "MINUTA DA ATA DE REGISTRO DE PREÇOS" in texto or "MINUTA DE ATA DE REGISTRO DE PREÇOS" in texto:
            if p._element is not arp_element:
                xml_str = etree.tostring(p._element, encoding='unicode').upper()
                if 'TOC' in xml_str or 'PAGEREF' in xml_str or 'HYPERLINK' in xml_str:
                    try:
                        parent = p._element.getparent()
                        if parent is not None:
                            parent.remove(p._element)
                            if len(parent.findall(qn('w:p'))) == 0:
                                parent.append(OxmlElement('w:p'))
                    except Exception:
                        pass
                        
    if arp_element is None:
        return
        
    el = arp_element
    while el is not None and el.getparent() is not body:
        el = el.getparent()
        
    if el is None:
        return
        
    try:
        idx_corte = elementos_body.index(el)
    except ValueError:
        return
        
    temp_idx = idx_corte
    limite_busca = max(0, idx_corte - 10)
    for j in range(idx_corte - 1, limite_busca - 1, -1):
        el_j = elementos_body[j]
        txt = re.sub(r'\s+', ' ', "".join(el_j.itertext()).strip().upper())
        try:
            xml_str = etree.tostring(el_j, encoding='unicode').upper()
        except Exception:
            xml_str = ""
            
        is_page_break = ('TYPE="PAGE"' in xml_str.replace(" ", "") or
                         "PAGEBREAKBEFORE" in xml_str.replace(" ", "") or
                         "BREAK" in xml_str)
                         
        if is_page_break:
            temp_idx = j
            break
            
        if txt.startswith("ANEXO") and len(txt) < 80:
            temp_idx = j
            
    idx_corte = temp_idx
    
    ultimo_elemento = elementos_body[-1] if elementos_body else None
    for el in elementos_body[idx_corte:]:
        if el is ultimo_elemento and el.tag.endswith('sectPr'):
            continue
        try:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        except Exception:
            pass

    if len(body.findall(qn('w:p'))) == 0:
        body.insert(0, OxmlElement('w:p'))

def preencher_documento(caminho_modelo: str, caminho_saida: str, dados: dict) -> str:
    doc = Document(caminho_modelo)
    e_arp = dados.get("E_ARP", False)
    remover_amostra = dados.get("__REMOVER_AMOSTRA__", False)
    remover_vistoria = dados.get("__REMOVER_VISTORIA__", False)
    apenas_contrato = dados.get("APENAS_CONTRATO", False) or dados.get("__APENAS_CONTRATO__", False)
    paragrafos_remover_set = set()
    
    for i, p in enumerate(doc.paragraphs):
        texto_upper = p.text.strip().upper()
        is_amostra = remover_amostra and "AMOSTRA" in texto_upper and len(texto_upper) < 60
        is_vistoria = remover_vistoria and "VISTORIA" in texto_upper and len(texto_upper) < 60
        if is_amostra or is_vistoria:
            paragrafos_remover_set.add(p)
            idx = i + 1
            while idx < len(doc.paragraphs):
                next_p = doc.paragraphs[idx]
                txt = next_p.text.strip()
                if not txt or (remover_amostra and "{{AMOSTRA}}" in txt) or (remover_vistoria and "{{VISTORIA}}" in txt):
                    paragrafos_remover_set.add(next_p)
                    idx += 1
                else:
                    break
            idx = i - 1
            while idx >= 0:
                prev_p = doc.paragraphs[idx]
                if not prev_p.text.strip():
                    paragrafos_remover_set.add(prev_p)
                    idx -= 1
                else:
                    break
                    
    for p in paragrafos_remover_set:
        try:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
                if len(parent.findall(qn('w:p'))) == 0:
                    parent.append(OxmlElement('w:p'))
        except Exception:
            pass
            
    PLACEHOLDERS_DOC = {"{{DFD}}", "{{ETP}}", "{{TR}}"}
    chaves_doc_processadas = set()
    
    for chave_ph in PLACEHOLDERS_DOC:
        caminho_arq = str(dados.get(chave_ph, "")).strip()
        if not caminho_arq:
            continue
            
        ext = os.path.splitext(caminho_arq)[1].lower()
        chaves_doc_processadas.add(chave_ph)
        
        for p in list(doc.paragraphs):
            if chave_ph in p.text:
                if not os.path.exists(caminho_arq):
                    _substituir_texto_mantendo_formatacao(p, chave_ph, f"[ERRO: ARQUIVO NÃO ENCONTRADO - {caminho_arq}]")
                elif ext == '.docx':
                    _mesclar_docx_no_local(doc, p, caminho_arq)
                else:
                    _substituir_texto_mantendo_formatacao(p, chave_ph, f"[ERRO: FORMATO NÃO SUPORTADO. ENVIE APENAS .DOCX - {os.path.basename(caminho_arq)}]")
                break
                
        dados = dict(dados)
        dados[chave_ph] = ""

    dados = dict(dados)
    tipo_dotacao = dados.get("{{TIPO_DOTACAO}}", "TEXTO")
    dotacao_b64_list = dados.get("{{DOTACAO_BASE64_LIST}}", [])

    if tipo_dotacao == "IMAGEM" and dotacao_b64_list:
        try:
            image_streams = []
            for b64_str in dotacao_b64_list:
                if b64_str:
                    image_data = base64.b64decode(b64_str)
                    image_streams.append(BytesIO(image_data))
            
            if image_streams:
                def substituir_dotacao_imagens(p):
                    if "{{DOTACAO}}" in p.text:
                        _substituir_texto_mantendo_formatacao(p, "{{DOTACAO}}", "")
                        run = p.add_run()
                        for img_stream in image_streams:
                            run.add_picture(img_stream, width=Inches(5.5))
                            img_stream.seek(0)
                
                for p in list(doc.paragraphs):
                    substituir_dotacao_imagens(p)
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in list(cell.paragraphs):
                                substituir_dotacao_imagens(p)
                                
                dados["{{DOTACAO}}"] = ""
        except Exception as e:
            pass

    dados.pop("{{DOTACAO_BASE64_LIST}}", None)
    dados.pop("{{TIPO_DOTACAO}}", None)

    if not e_arp:
        _remover_secao_arp(doc)
        
    for paragrafo in list(doc.paragraphs):
        _processar_paragrafo(paragrafo, dados, e_arp)
        
    for tabela in doc.tables:
        colunas_para_remover = []
        for i, coluna in enumerate(tabela.columns):
            remover_esta = False
            for cell in coluna.cells:
                if "__REMOVER_COLUNA__" in cell.text:
                    remover_esta = True
                    break
            if remover_esta:
                colunas_para_remover.append(i)
                
        if len(colunas_para_remover) == len(tabela.columns) and len(tabela.columns) > 0:
            try:
                tbl = tabela._tbl
                parent = tbl.getparent()
                if parent is not None:
                    parent.remove(tbl)
            except Exception:
                pass
        else:
            for col_idx in reversed(colunas_para_remover):
                for row in tabela.rows:
                    try:
                        tc = row.cells[col_idx]._tc
                        parent = tc.getparent()
                        if parent is not None:
                            parent.remove(tc)
                    except Exception:
                        pass
                        
        if tabela._tbl.getparent() is not None:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in list(celula.paragraphs):
                        _processar_paragrafo(paragrafo, dados, e_arp)
                    
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for paragrafo in list(header.paragraphs):
                    _processar_paragrafo(paragrafo, dados, e_arp)
                for tabela in header.tables:
                    for linha in tabela.rows:
                        for celula in linha.cells:
                            for paragrafo in list(celula.paragraphs):
                                _processar_paragrafo(paragrafo, dados, e_arp)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for paragrafo in list(footer.paragraphs):
                    _processar_paragrafo(paragrafo, dados, e_arp)
                for tabela in footer.tables:
                    for linha in tabela.rows:
                        for celula in linha.cells:
                            for paragrafo in list(celula.paragraphs):
                                _processar_paragrafo(paragrafo, dados, e_arp)
                                
    def limpar_realce_verde(elemento_raiz):
        for node in elemento_raiz.xpath('.//w:highlight'):
            val = node.get(qn('w:val'))
            if val in ['green', 'brightGreen']:
                parent = node.getparent()
                if parent is not None:
                    parent.remove(node)
                    
    limpar_realce_verde(doc.element)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                limpar_realce_verde(header._element)
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                limpar_realce_verde(footer._element)
                
    doc.save(caminho_saida)
    remover_paginas_em_branco(caminho_saida)
    
    if apenas_contrato:
        _remover_ultimo_sectpr(caminho_saida)
        
    return caminho_saida

def _substituir_texto_mantendo_formatacao(paragrafo, marcador, valor_str):
    if marcador not in paragrafo.text:
        return False
    texto_p = paragrafo.text
    idx_inicio = texto_p.find(marcador)
    while idx_inicio != -1:
        idx_fim = idx_inicio + len(marcador)
        pos_atual = 0
        runs_envolvidos = []
        for run in paragrafo.runs:
            texto_run = run.text
            len_run = len(texto_run)
            pos_fim_run = pos_atual + len_run
            if pos_fim_run > idx_inicio and pos_atual < idx_fim:
                runs_envolvidos.append({
                    'run': run,
                    'inicio_run': pos_atual,
                    'fim_run': pos_fim_run,
                    'texto_original': texto_run
                })
            pos_atual = pos_fim_run
        if runs_envolvidos:
            for j, info in enumerate(runs_envolvidos):
                run = info['run']
                txt = info['texto_original']
                inicio_intersecao = max(0, idx_inicio - info['inicio_run'])
                fim_intersecao = min(len(txt), idx_fim - info['inicio_run'])
                prefixo = txt[:inicio_intersecao]
                sufixo = txt[fim_intersecao:]
                if j == 0:
                    run.text = prefixo + valor_str + sufixo
                else:
                    run.text = prefixo + sufixo
        texto_p = paragrafo.text
        idx_inicio = texto_p.find(marcador)
    return True

def _aplicar_formatacao_markdown_avancado(paragrafo):
    if "***" not in paragrafo.text and "**" not in paragrafo.text:
        return
    runs = list(paragrafo.runs)
    for run in runs:
        texto_run = run.text
        if "***" in texto_run or "**" in texto_run:
            parts = re.split(r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*)', texto_run, flags=re.DOTALL)
            if len(parts) > 1:
                run.text = ""
                run_base_element = run._element
                for part in parts:
                    if not part:
                        continue
                    new_run = paragrafo.add_run()
                    new_run._element.getparent().remove(new_run._element)
                    run_base_element.addnext(new_run._element)
                    run_base_element = new_run._element
                    if run._element.rPr is not None:
                        new_run._element.insert(0, copy.deepcopy(run._element.rPr))
                    if part.startswith('***') and part.endswith('***'):
                        new_run.text = part[3:-3]
                        new_run.font.bold = True
                        new_run.font.italic = True
                        new_run.font.underline = True
                    elif part.startswith('**') and part.endswith('**'):
                        new_run.text = part[2:-2]
                        new_run.font.bold = True
                    else:
                        new_run.text = part
            try:
                run._element.getparent().remove(run._element)
            except Exception:
                pass

def _inserir_multilinhas_safe(paragrafo, marcador, valor_str):
    texto_p = paragrafo.text
    idx_inicio = texto_p.find(marcador)
    rPr_base = None
    if idx_inicio != -1:
        pos_atual = 0
        for run in paragrafo.runs:
            len_run = len(run.text)
            if pos_atual + len_run > idx_inicio:
                if run._element.rPr is not None:
                    rPr_base = copy.deepcopy(run._element.rPr)
                break
            pos_atual += len_run
    if rPr_base is None and paragrafo.runs:
        for r in paragrafo.runs:
            if r._element.rPr is not None:
                rPr_base = copy.deepcopy(r._element.rPr)
                break
    valor_str = str(valor_str).replace('\r', '')
    valor_str = re.sub(r'\n{3,}', '\n\n', valor_str)
    linhas = valor_str.strip('\n').split('\n')
    if not linhas or (len(linhas) == 1 and linhas[0] == ""):
        _substituir_texto_mantendo_formatacao(paragrafo, marcador, "")
        return []
    _substituir_texto_mantendo_formatacao(paragrafo, marcador, linhas[0])
    paragrafo_atual = paragrafo
    novos_paragrafos = []
    for linha in linhas[1:]:
        novo_p = OxmlElement("w:p")
        paragrafo_atual._p.addnext(novo_p)
        if paragrafo_atual._p.pPr is not None:
            novo_p.append(copy.deepcopy(paragrafo_atual._p.pPr))
        novo_paragrafo_obj = Paragraph(novo_p, paragrafo_atual._parent)
        novo_run = novo_paragrafo_obj.add_run(linha)
        if rPr_base is not None:
            novo_run._element.insert(0, copy.deepcopy(rPr_base))
        paragrafo_atual = novo_paragrafo_obj
        novos_paragrafos.append(novo_paragrafo_obj)
    return novos_paragrafos

def _processar_paragrafo(paragrafo, dados, e_arp):
    apagou_algo = False
    for run in paragrafo.runs:
        is_green = False
        highlight_element = None
        if run._element.rPr is not None:
            highlight = run._element.rPr.find(qn('w:highlight'))
            if highlight is not None:
                val = highlight.get(qn('w:val'))
                if val in ['green', 'brightGreen']:
                    is_green = True
                    highlight_element = highlight
        if is_green:
            if not e_arp:
                if run.text:
                    run.text = ""
                    apagou_algo = True
            else:
                if highlight_element is not None:
                    run._element.rPr.remove(highlight_element)
                    
    paragrafos_para_processar = [paragrafo]
    for chave, valor in dados.items():
        if chave in ["E_ARP", "__REMOVER_AMOSTRA__", "__REMOVER_VISTORIA__", "APENAS_CONTRATO", "__APENAS_CONTRATO__"]:
            continue
        marcador = chave if chave.startswith("{{") and chave.endswith("}}") else f"{{{{{chave}}}}}"
        if isinstance(valor, list):
            valor_str = "\n".join([str(v) for v in valor if str(v).strip()])
        else:
            valor_str = str(valor)
        novos_adicionados = []
        for p_atual in list(paragrafos_para_processar):
            if marcador in p_atual.text:
                if '\n' in valor_str:
                    novos_ps = _inserir_multilinhas_safe(p_atual, marcador, valor_str)
                    novos_adicionados.extend(novos_ps)
                    if not valor_str.strip() and not p_atual.text.strip():
                        apagou_algo = True
                else:
                    _substituir_texto_mantendo_formatacao(p_atual, marcador, valor_str)
                    if not valor_str.strip() and not p_atual.text.strip():
                        apagou_algo = True
        paragrafos_para_processar.extend(novos_adicionados)
        
    for p_atual in paragrafos_para_processar:
        _aplicar_formatacao_markdown_avancado(p_atual)
        if apagou_algo and not p_atual.text.strip():
            try:
                p_element = p_atual._element
                parent = p_element.getparent()
                if parent is not None:
                    parent.remove(p_element)
                    if len(parent.findall(qn('w:p'))) == 0:
                        parent.append(OxmlElement('w:p'))
            except Exception:
                pass