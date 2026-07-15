import os
import re
import json
from copy import deepcopy
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.oxml.ns import qn

PLACEHOLDER_RE = re.compile(r"\{\{\s*[^{}]+?\s*\}\}")


def _safe_color_rgb(run):
    try:
        color = run.font.color
        if color.type is not None:
            return color.rgb
    except Exception:
        pass
    return None


def extrair_placeholders_documento(caminho_docx):
    doc = Document(caminho_docx)
    placeholders = set()

    def coletar(texto):
        placeholders.update(PLACEHOLDER_RE.findall(texto or ""))

    def iterar_blocos(blocos):
        for b in blocos:
            for p in getattr(b, "paragraphs", []):
                coletar(p.text)
            for t in getattr(b, "tables", []):
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            coletar(p.text)

    iterar_blocos([doc])
    for section in doc.sections:
        iterar_blocos([section.header, section.footer])

    return placeholders


def extrair_placeholders_modelos(pasta_modelos, arquivos_base):
    placeholders = set()
    for arquivo in arquivos_base:
        caminho = os.path.join(pasta_modelos, arquivo)
        if os.path.exists(caminho):
            placeholders.update(extrair_placeholders_documento(caminho))
    return placeholders


def _copy_numbering_properties(source_paragraph, target_paragraph):
    if source_paragraph._p.pPr is None or source_paragraph._p.pPr.numPr is None:
        return
    target_pPr = target_paragraph._p.get_or_add_pPr()
    for child in list(target_pPr):
        if child.tag == qn("w:numPr"):
            target_pPr.remove(child)
    target_pPr.append(deepcopy(source_paragraph._p.pPr.numPr))


def _split_linear_content_for_paragraphs(linear_content):
    paragraphs = [[]]
    for segment in linear_content:
        if "\n" in segment["text"]:
            parts = segment["text"].split("\n")
            for i, part in enumerate(parts):
                paragraphs[-1].append({**segment, "text": part})
                if i < len(parts) - 1:
                    paragraphs.append([])
        else:
            paragraphs[-1].append(segment)
    return paragraphs


def _inserir_tabela(paragraph, json_str):
    try:
        dados = json.loads(json_str)
        if not dados:
            return
        parent = paragraph._parent
        table = parent.add_table(rows=1, cols=len(dados[0]), width=Inches(6.0))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, key in enumerate(dados[0].keys()):
            hdr_cells[i].text = str(key)
        for item in dados:
            row_cells = table.add_row().cells
            for i, key in enumerate(item.keys()):
                row_cells[i].text = str(item[key])
        paragraph._p.addnext(table._tbl)
    except Exception as e:
        new_run = paragraph.add_run(f"[ERRO AO GERAR TABELA: {e}]")
        new_run.font.color.rgb = RGBColor(255, 0, 0)
        new_run.bold = True


def _apply_segments_to_paragraph(paragraph, segments, extracted_runs_data):
    alinhamento_original = paragraph.alignment
    paragraph.clear()
    paragraph.alignment = alinhamento_original
    for segment in segments:
        if not segment.get("text"):
            continue

        original_run_data = extracted_runs_data[segment["original_run_index"]]

        parts = segment["text"].split("\n")
        for i, part in enumerate(parts):
            if part:
                if part.startswith("__IMG__"):
                    img_path = part.replace("__IMG__", "")
                    if os.path.exists(img_path):
                        new_run = paragraph.add_run()
                        new_run.add_picture(img_path, width=Inches(6.0))
                    else:
                        new_run = paragraph.add_run(f"[IMAGEM NÃO ENCONTRADA: {img_path}]")
                        new_run.font.color.rgb = RGBColor(255, 0, 0)
                        new_run.bold = True
                elif part.startswith("__TABLE__"):
                    json_str = part.replace("__TABLE__", "")
                    _inserir_tabela(paragraph, json_str)
                else:
                    new_run = paragraph.add_run(part)
                    new_run.style = original_run_data["style"]
                    new_run.bold = original_run_data["bold"]
                    new_run.italic = original_run_data["italic"]
                    new_run.underline = original_run_data["underline"]
                    if original_run_data["font_name"]:
                        new_run.font.name = original_run_data["font_name"]
                    if original_run_data["font_size"]:
                        new_run.font.size = original_run_data["font_size"]
                    if original_run_data["color_rgb"] is not None:
                        new_run.font.color.rgb = original_run_data["color_rgb"]
            if i < len(parts) - 1:
                paragraph.add_run().add_break()


def replace_text_in_paragraph(paragraph, replacements):
    extracted_runs_data = []
    for run in paragraph.runs:
        extracted_runs_data.append({
            "text": run.text or "",
            "style": run.style,
            "bold": run.bold,
            "italic": run.italic,
            "underline": run.underline,
            "font_name": run.font.name,
            "font_size": run.font.size,
            "color_rgb": _safe_color_rgb(run),
        })

    linear_content = []
    for i, run_data in enumerate(extracted_runs_data):
        linear_content.append({"text": run_data["text"], "original_run_index": i})

    texto_completo = "".join(rd["text"] for rd in extracted_runs_data)
    precisa_reconstruir = any(
        old_text in texto_completo
        and not any(old_text in segment["text"] for segment in linear_content)
        for old_text in replacements
    )

    if precisa_reconstruir and extracted_runs_data:
        for old_text, new_text in replacements.items():
            texto_completo = texto_completo.replace(old_text, str(new_text))
        linear_content = [{"text": texto_completo, "original_run_index": 0}]
    else:
        for old_text, new_text in replacements.items():
            new_linear_content = []
            for segment in linear_content:
                if old_text in segment["text"]:
                    parts = segment["text"].split(old_text)
                    for i, part in enumerate(parts):
                        if part:
                            new_linear_content.append({**segment, "text": part})
                        if i < len(parts) - 1:
                            new_linear_content.append({**segment, "text": str(new_text)})
                else:
                    new_linear_content.append(segment)
            linear_content = new_linear_content

    list_paragraph = paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None
    paragraph_groups = _split_linear_content_for_paragraphs(linear_content) if list_paragraph else [linear_content]

    if len(paragraph_groups) == 1:
        _apply_segments_to_paragraph(paragraph, paragraph_groups[0], extracted_runs_data)
        return

    for segments in reversed(paragraph_groups[:-1]):
        new_para = paragraph.insert_paragraph_before(text=None, style=paragraph.style)
        _copy_numbering_properties(paragraph, new_para)
        _apply_segments_to_paragraph(new_para, segments, extracted_runs_data)

    _apply_segments_to_paragraph(paragraph, paragraph_groups[-1], extracted_runs_data)


def modificar_documento(cam_origem, cam_destino, modificacoes):
    doc = Document(cam_origem)

    def substituir_blocos(blocos):
        for b in blocos:
            for p in getattr(b, "paragraphs", []):
                replace_text_in_paragraph(p, modificacoes)
            for t in getattr(b, "tables", []):
                for r in t.rows:
                    for c in r.cells:
                        for p in c.paragraphs:
                            replace_text_in_paragraph(p, modificacoes)

    substituir_blocos([doc])
    for section in doc.sections:
        substituir_blocos([section.header, section.footer])

    doc.save(cam_destino)
